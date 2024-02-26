# main_window.py
# This file creates the main window for the GUI.

from customtkinter import *
from tkinter import filedialog, messagebox, PhotoImage
from PIL import Image, ImageTk
from app_logging import logger
from gui.visualization_panel import VisualizationPanel
from model.neural_network import NeuralNetworkModel
from data.preprocessing import preprocess_data
from docx import Document

# Global list of feature names for use in plotting
feature_names = ['N', 'f', 'V', 'T']

class MainWindow(CTk):
    def __init__(self):
        super().__init__()
        logger.info("Starting MainWindow initialization")
        
        self.title('Ultra_Aim_Pro')
        self.minsize(300, 200)  # Adjust the size to fit your needs
        set_appearance_mode("dark")  # Dark mode for the entire application
                
        # Create the Tabview widget for a tabbed interface
        self.tabview = CTkTabview(master=self)
        self.tabview.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # Create the 'Main' tab and 'Analyze' tab
        self.main_tab = self.tabview.add("Main")
        self.visualization_tab = self.tabview.add("Analyze")

        # Initialize the visualization panel
        self.visualization_panel = VisualizationPanel(self.visualization_tab)
        self.visualization_panel.grid(row=0, column=0, sticky="nsew")
        self.visualization_tab.grid_rowconfigure(0, weight=1)
        self.visualization_tab.grid_columnconfigure(0, weight=1)

        
        self.download_button = CTkButton(self.visualization_tab, text='Download', command=self.download_for_matlab)
        self.download_button.grid(row=1, column=0, pady=10, padx=10)

        self.download_image_button = CTkButton(self.visualization_tab, text='Download as Image', command=self.download_as_image)
        self.download_image_button.grid(row=1, column=1, pady=10, padx=10)

        
        self.download_analysis_button = CTkButton(self.visualization_tab, text='Download Analysis', command=self.download_analysis)
        self.download_analysis_button.grid(row=2, column=0, pady=10, padx=10)

        # Create widgets in the 'Main' tab
        self.create_main_tab_widgets()

        # Placeholder for data and model
        self.data = None
        self.model = None
                
        logger.info("MainWindow initialized successfully.")

        
    def save_analysis_to_doc(self, analysis, file_path):
        # Create a new Document
        doc = Document()
        doc.add_heading('Graph Analysis', 0)

        # Add analysis for each graph
        for graph_name, graph_analysis in analysis.items():
            doc.add_heading(graph_name, level=1)
            doc.add_paragraph(graph_analysis)
        
        # Save the document
        doc.save(file_path)

    def download_analysis(self):
        # This function will be called when the 'Download' button is clicked
        file_path = filedialog.asksaveasfilename(
            title='Save Graph Analysis',
            filetypes=[('Word Documents', '*.docx')],
            defaultextension='.docx'
        )
        if file_path:
            # Get the analysis data (this is just an example, you will need to provide the actual content)
            analysis = {
                'Training & Validation Loss': (
                    "The training and validation loss graphs show how the model's error decreases over each epoch. "
                    "The sharp decline in training loss suggests that the model quickly learns to fit the training data. "
                    "However, the validation loss plateau indicates that improvements in the model's performance on the training data "
                    "do not translate to equivalent improvements on unseen data. This discrepancy can signal overfitting, "
                    "where the model learns the training data too well, including its noise and outliers, "
                    "which does not generalize well to new data."
                ),
                'Prediction Accuracy': (
                    "The prediction accuracy graph shows a scatter plot of the model's predicted reliability against the true reliability values. "
                    "The cluster of points around the dashed line of perfect prediction indicates that the model has a good level of predictive accuracy. "
                    "However, there are deviations, especially in higher reliability values, where the model tends to underestimate the reliability. "
                    "This could suggest that the model may need further tuning to handle higher reliability ranges or that there are factors affecting high reliability that the model is not considering."
                ),
                'Impact of N': (
                    "The graph for the impact of 'N' on reliability shows a spread of predicted versus actual reliability which does not indicate a clear trend. "
                    "This could imply that 'N' is not a strong predictor for reliability on its own, or that its relationship with reliability is non-linear and possibly influenced by other variables. "
                    "A more complex model or feature engineering may be needed to capture the true impact of 'N' on reliability."
                ),
                'Impact of V': (
                    "The 'Impact of V' graph presents a vertical clustering of points at specific 'V' values, suggesting a categorical or discrete nature of 'V'. "
                    "The overlaps of predicted and actual reliability values indicate that the model can capture the effect of 'V' on reliability to some extent. "
                    "However, there is notable variance in the predictions at the extreme values of 'V', which may require additional investigation."
                ),
                'Impact of f': (
                    "In the 'Impact of f' graph, the distribution of points shows a random pattern, indicating a weak or complex relationship between 'f' and reliability. "
                    "The model's predictions do not consistently align with the actual values across the range of 'f'. "
                    "This suggests that 'f' may not be a significant predictor in the current model's form, or it interacts with reliability in a way that the model is currently not capturing."
                ),
                'Impact of T': (
                    "The 'Impact of T' graph depicts a horizontal banding pattern, with the predicted values generally matching the actual reliability across different values of 'T'. "
                    "This indicates that while 'T' varies, its impact on reliability is not strongly captured by the model. "
                    "If 'T' is an important factor in theory, the model may need to be reassessed to ensure it can leverage 'T' effectively for prediction."
                ),
            }

            self.save_analysis_to_doc(analysis, file_path)
    

    def download_for_matlab(self):
        file_path = filedialog.asksaveasfilename(
            title='Save Graphs for MATLAB',
            filetypes=[('MATLAB Files', '*.mat')],
            defaultextension='.mat'
        )
        if file_path:
            self.visualization_panel.save_graphs_for_matlab(file_path)


    def download_as_image(self):
        # Ask for the directory instead of the file path since we are saving multiple files
        directory = filedialog.askdirectory(title='Select Directory to Save Images')
        if directory:
            for idx, fig in enumerate(self.visualization_panel.figures):
                # Construct a unique filename for each figure
                file_path = f"{directory}/figure_{idx + 1}.png"
                self.visualization_panel.save_graphs_as_image(file_path, fig)

        
    
    def create_main_tab_widgets(self):
        """Create widgets for the main tab, ensuring the logo image is correctly displayed."""
                        
        # Logo (make sure to use the correct path for your logo)
        logo_path = r'G:\My Drive\Final Project Boris_E\final_project\Ultra_Aim_Pro - Final\Final Project\graphics\logo.png'
        logo_image = Image.open(logo_path)
        self.logo_image = ImageTk.PhotoImage(logo_image)  # Convert to PhotoImage
        self.logo_label = CTkLabel(self.main_tab, image=self.logo_image, text="")
        self.logo_label.grid(row=0, column=0, columnspan=3, pady=10, padx=10)



        # Button to upload the data file
        upload_img = Image.open(r'G:\My Drive\Final Project Boris_E\final_project\Ultra_Aim_Pro - Final\Final Project\graphics\upload.png')
        self.upload_button = CTkButton(self.main_tab, text='Upload Data', command=self.upload_data,
                                           image=CTkImage(dark_image=upload_img, light_image=upload_img),
                                           corner_radius=10)
        self.upload_button.grid(row=1, column=0, columnspan=3, pady=10, padx=10)

        # Label to show the file info
        self.file_info_label = CTkLabel(self.main_tab, text='No file selected', text_color="white")
        self.file_info_label.grid(row=2, column=0, columnspan=3, pady=10, padx=10)


        # Button to start the training
        train_img = Image.open(r'G:\My Drive\Final Project Boris_E\final_project\Ultra_Aim_Pro - Final\Final Project\graphics\configuration.png')
        self.train_button = CTkButton(self.main_tab, text='Train Model', state='disabled', command=self.train_model,
                                          image=CTkImage(dark_image=train_img, light_image=train_img),
                                          fg_color="#333333", hover_color="#444444", text_color="white", corner_radius=10)
        self.train_button.grid(row=3, column=0, columnspan=3, pady=10, padx=10)


        # File info label, progress label, save and load model buttons don't need adjustments for image management
        self.file_info_label = CTkLabel(self.main_tab, text='No file selected', text_color="white")
        self.file_info_label.grid(row=2, column=0, columnspan=3, pady=10, padx=10)

        self.progress_label = CTkLabel(self.main_tab, text='Ready', text_color="white")
        self.progress_label.grid(row=6, column=4, columnspan=3, pady=10, padx=10)

        self.save_model_button = CTkButton(self.main_tab, text='Save Model', command=self.save_model)
        self.save_model_button.grid(row=8, column=4, pady=10, padx=10)

        self.load_model_button = CTkButton(self.main_tab, text='Load Model', command=self.load_model)
        self.load_model_button.grid(row=10, column=4, pady=10, padx=10)

        # Create the sliders with entry boxes in the 'Main' tab
        self.create_sliders_in_main_tab()

    
    def download_for_matlab(self):
        file_path = filedialog.asksaveasfilename(
            title='Save Graphs for MATLAB',
            filetypes=[('MATLAB Files', '*.mat')],
            defaultextension='.mat'
        )
        if file_path:
            self.visualization_panel.save_graphs_for_matlab(file_path)


    
    
    def create_sliders_in_main_tab(self):
        # Define the parameters for the sliders
        slider_params = [
            {'row': 5, 'label_text': 'Dense Layer 1 Units:', 'from_': 16, 'to': 128, 'default_value': 64, 'number_of_steps': 113},
            {'row': 7, 'label_text': 'Dense Layer 2 Units:', 'from_': 16, 'to': 128, 'default_value': 32, 'number_of_steps': 113},
            {'row': 9, 'label_text': 'Learning Rate:', 'from_': 0.0001, 'to': 0.01, 'default_value': 0.001, 'number_of_steps': 1001},
            {'row': 11, 'label_text': 'Validation Split:', 'from_': 0.01, 'to': 0.2, 'default_value': 0.07, 'number_of_steps': 20},
            {'row': 13, 'label_text': 'Epochs:', 'from_': 100, 'to': 2000, 'default_value': 1000, 'number_of_steps': 20},
            {'row': 15, 'label_text': 'Batch Size:', 'from_': 16, 'to': 128, 'default_value': 77, 'number_of_steps': 113}
        ]

        # Create the sliders with entry boxes
        self.sliders = {}
        for params in slider_params:
            label = CTkLabel(self.main_tab, text=params['label_text'], text_color="white", fg_color="#333333")
            label.grid(row=params['row'], column=0, pady=10, padx=10, sticky='e')

            entry = CTkEntry(self.main_tab, width=120, fg_color="#333333", text_color="white")
            entry.grid(row=params['row'], column=1, pady=15, padx=15, sticky='w')
            entry.insert(0, str(params['default_value']))  # Set the default value

            slider = CTkSlider(self.main_tab, from_=params['from_'], to=params['to'], number_of_steps=params['number_of_steps'],
                                   fg_color="#333333", button_color="#555555", button_hover_color="#666666")
            
            slider.set(params['default_value'])  # Set the default value
            slider.grid(row=params['row'] + 1, column=0, columnspan=3, pady=10, padx=10)

            label, entry, slider = self.create_slider_with_entry(params)
            self.sliders[params['label_text']] = {'label': label, 'entry': entry, 'slider': slider}

    def create_slider_with_entry(self, params):
        # Label for the slider
        slider_label = CTkLabel(self.main_tab, text=params['label_text'])
        slider_label.grid(row=params['row'], column=0, pady=10, padx=10, sticky='e')

        # Entry for the slider value
        slider_value_entry = CTkEntry(self.main_tab, width=120)
        slider_value_entry.grid(row=params['row'], column=1, pady=15, padx=15, sticky='w')
        slider_value_entry.insert(0, str(params['default_value']))  # Set the default value

        # Slider
        slider = CTkSlider(self.main_tab, from_=params['from_'], to=params['to'], number_of_steps=params['number_of_steps'])
        slider.set(params['default_value'])  # Set the default value
        slider.grid(row=params['row'] + 1, column=0, columnspan=3, pady=10, padx=10)

        # Bind the slider movement to update the entry
        def slider_moved(event):
            slider_value_entry.delete(0, 'end')
            slider_value_entry.insert(0, str(slider.get()))

        slider.bind('<B1-Motion>', slider_moved)
        slider.bind('<ButtonRelease-1>', slider_moved)

        # Bind the entry to update the slider
        def entry_changed(event):
            try:
                value = float(slider_value_entry.get())
                if params['from_'] <= value <= params['to']:
                    slider.set(value)
            except ValueError:
                pass  # If the entry is not a valid number, do nothing

        slider_value_entry.bind('<Return>', entry_changed)

        return slider_label, slider_value_entry, slider

    def upload_data(self):
        file_path = filedialog.askopenfilename(
            title='Select Dataset',
            filetypes=[('Excel Files', '*.xlsx'), ('All Files', '*.*')]
        )
        print(f"File path chosen: {file_path}")  # Debug print
        if file_path:
            try:
                # Update the file info label with the selected file
                self.file_info_label.configure(text=f'Selected: {file_path}')
                logger.info(f"ℹ️ File path chosen: {file_path}")  # Replace print with logging

                # Preprocess the data
                self.data = preprocess_data(file_path)
                logger.info("ℹ️ Data preprocessed successfully.")
                print("Data preprocessed successfully.")  # Debug print
                # Enable the train button
                self.train_button.configure(state='normal')
            except Exception as e:
                messagebox.showerror('Error', f'An error occurred while loading the data: {e}')
                logger.critical(f"⛔ An error occurred while loading the data: {e}")  # Replace messagebox with logging
                print(f"An error occurred: {e}")  # Debug print
    
    
    def train_model(self):
            # Extract values from sliders
            dense1_units = int(self.sliders['Dense Layer 1 Units:']['entry'].get())
            dense2_units = int(self.sliders['Dense Layer 2 Units:']['entry'].get())
            learning_rate = float(self.sliders['Learning Rate:']['entry'].get())
            validation_split = float(self.sliders['Validation Split:']['entry'].get())
            epochs = int(float(self.sliders['Epochs:']['entry'].get()))
            batch_size = int(self.sliders['Batch Size:']['entry'].get())

            if self.data:
                # Unpack preprocessed data
                X_train, X_test, y_train, y_test, _ = self.data
                self.model = NeuralNetworkModel(input_shape=(X_train.shape[1],), dense1_units=dense1_units, dense2_units=dense2_units, learning_rate=learning_rate)

                # Train the model
                history = self.model.train(X_train, y_train, validation_split, epochs, batch_size)

                # Make predictions
                predictions = self.model.predict(X_test)

                # Plot results using the visualization panel
                self.visualization_panel.plot_loss(history)
                self.visualization_panel.plot_predictions(y_test, predictions)
                self.visualization_panel.plot_parameter_impact(X_test, y_test, predictions, feature_names)

                # Evaluate the model
                r2, mse = self.model.evaluate(X_test, y_test)
                r2_score_percent = r2 * 100

                # Show evaluation results
                logger.info(f"ℹ️ R^2 Score: {r2_score_percent:.4f}%\nMSE: {mse:.4f}")
                messagebox.showinfo('Model Evaluation', f'R^2 Score: {r2_score_percent:.4f}%\nMSE: {mse:.4f}')

                # Update progress label
                self.progress_label.configure(text='Training completed.')
            else:
                # Warn if no data is loaded
                messagebox.showwarning('Warning', 'Please upload data before training.')
                logger.warning("⚠️ Please upload data before training.")
        
        
    def save_model(self):
        file_path = filedialog.asksaveasfilename(
            title='Save Model',
            filetypes=[('H5 Files', '*.h5'), ('All Files', '*.*')],
            defaultextension='.h5'
        )
        if file_path:
            self.model.save_model(file_path)

    def load_model(self):
        file_path = filedialog.askopenfilename(
            title='Load Model',
            filetypes=[('H5 Files', '*.h5'), ('All Files', '*.*')]
        )
        if file_path:
            self.model = NeuralNetworkModel.load_model(file_path)
    

if __name__ == '__main__':
    logger.info("Starting application")
    root = MainWindow()
    root.mainloop()
    logger.info("Application closed")