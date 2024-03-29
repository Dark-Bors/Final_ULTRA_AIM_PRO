# visualization_panel.py
# This file defines the panel for data visualization.

import tkinter as tk
import scipy.io
import matplotlib.pyplot as plt
import numpy as np
from tkinter import Scrollbar, Frame, filedialog
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
from app_logging import logger
from docx import Document
from scipy.stats import linregress


class VisualizationPanel(tk.Frame):
    def __init__(self, master):
        super().__init__(master)
        self.config(borderwidth=2, relief='sunken')
        logger.info("ℹ️ Initializing the VisualizationPanel.")
        self.master = master
        
        # Set grid weight configurations
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # Create a canvas with a scrollbar
        self.canvas = tk.Canvas(self)
        self.scrollbar = Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = Frame(self.canvas)

        # Configure the canvas
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        # Bind the canvas configuration to an event
        self.scrollable_frame.bind("<Configure>", self.on_frame_configure)
        self.canvas.bind("<Configure>", self.on_canvas_configure)

        # Place the widgets using grid
        self.scrollbar.grid(row=0, column=1, sticky="ns")
        self.canvas.grid(row=0, column=0, sticky="nsew")

        # Initialize variables to store figures and canvases
        self.figures = []
        self.canvases = []

    #----------------new part : ---------------
    
    def download_analysis(self):
        # This method will handle the button click event to download the analysis
        file_path = filedialog.asksaveasfilename(
            title='Save Analysis Document',
            filetypes=[('Word Documents', '*.docx')],
            defaultextension='.docx'
        )
        if file_path:  # If a file path was selected
            self.save_analysis_to_doc(file_path)  # Call the method with the file_path


    def save_analysis_to_doc(self, file_path):        # Create a new Document
        doc = Document()
        doc.add_heading('Graph Analysis', 0)

        # Analysis for the Training & Validation Loss plot
        train_loss_data = self.figures[0].axes[0].lines[0].get_ydata()
        val_loss_data = self.figures[0].axes[0].lines[1].get_ydata()
        epochs = range(len(train_loss_data))  # Assuming epochs are along the x-axis
        analysis_text = self.analyze_loss_graph(train_loss_data, val_loss_data, epochs)
        doc.add_heading(f'Figure 1: Training & Validation Loss', level=1)
        doc.add_paragraph(analysis_text)

        # Analysis for the Prediction Accuracy plot
        predicted_reliability_data = self.figures[1].axes[0].collections[0].get_offsets().data[:, 1]
        true_values_data = self.figures[1].axes[0].collections[0].get_offsets().data[:, 0]
        analysis_text = self.analyze_prediction_accuracy(true_values_data, predicted_reliability_data)
        doc.add_heading(f'Figure 2: Prediction Accuracy', level=1)
        doc.add_paragraph(analysis_text)

        # Analysis for the Impact of Features (N, V, f, T) plots
        for i, feature_name in enumerate(['N', 'V', 'f', 'T']):
            feature_data = self.figures[2].axes[i].collections[0].get_offsets().data[:, 0]
            reliability_data = self.figures[2].axes[i].collections[0].get_offsets().data[:, 1]
            analysis_text = self.analyze_feature_impact(feature_data, reliability_data, feature_name)
            doc.add_heading(f'Figure 3: Impact of {feature_name}', level=1)
            doc.add_paragraph(analysis_text)

        # Save the document
        doc.save(file_path)
                
    #------------end of new part ! ------------
    
    def on_frame_configure(self, event=None):
        # Reset the scroll region to encompass the inner frame
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_canvas_configure(self, event):
        # Resize the inner frame to match the canvas
        self.canvas.itemconfig("self.scrollable_frame", width=event.width)

        # Update the canvas scrolling region
        self.on_frame_configure()

    def plot_loss(self, history):
        fig = Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        ax.clear()
        ax.plot(history.history['loss'], label='Train Loss')
        ax.plot(history.history['val_loss'], label='Validation Loss')
        ax.set_title('Training & Validation Loss')
        ax.set_ylabel('Loss')
        ax.set_xlabel('Epoch')
        ax.legend(loc='upper right')
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.scrollable_frame)
        canvas.draw()
        widget = canvas.get_tk_widget()
        widget.pack(fill=tk.BOTH, expand=True)

        # Store the figure and canvas
        self.figures.append(fig)
        self.canvases.append(canvas)

    def plot_predictions(self, y_test, predicted_reliability):
        fig = Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        ax.clear()
        ax.scatter(y_test, predicted_reliability, label='Predicted Reliability')
        ax.plot([y_test.min(), y_test.max()], [y_test.min(), y_test.max()], 'k--', lw=4, label='Perfect Prediction')
        ax.set_xlabel('True Values [Reliability]')
        ax.set_ylabel('Predictions [Reliability]')
        ax.set_title('Prediction Accuracy')
        ax.legend(loc='upper left')
        ax.grid(True)
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.scrollable_frame)
        canvas.draw()
        widget = canvas.get_tk_widget()
        widget.pack(fill=tk.BOTH, expand=True)

        # Store the figure and canvas
        self.figures.append(fig)
        self.canvases.append(canvas)

    def plot_parameter_impact(self, X_test, y_test, predicted_reliability, feature_names):
        # Determine the layout of the subplots
        num_features = len(feature_names)
        num_cols = 2
        num_rows = int(np.ceil(num_features / num_cols))
        fig, axs = plt.subplots(num_rows, num_cols, figsize=(12, num_rows * 4))

        for i, feature_name in enumerate(feature_names):
            ax = axs[i // num_cols, i % num_cols]
            ax.scatter(X_test[:, i], predicted_reliability, label='Predicted Reliability')
            ax.scatter(X_test[:, i], y_test, color='red', label='Actual Reliability', alpha=0.5)
            ax.set_xlabel(feature_name)
            ax.set_ylabel('Reliability')
            ax.set_title(f'Impact of {feature_name}')
            ax.legend()
            ax.grid(True)

        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.scrollable_frame)
        canvas.draw()
        widget = canvas.get_tk_widget()
        widget.pack(fill=tk.BOTH, expand=True)

        # Store the figure and canvas
        self.figures.append(fig)
        self.canvases.append(canvas)

    def clear_plots(self):
        logger.info("ℹ️ Clearing plots.")
        for canvas in self.canvases:
            canvas.get_tk_widget().pack_forget()
            canvas.get_tk_widget().destroy()
        self.figures.clear()
        self.canvases.clear()
        logger.debug("🐛 Plots cleared from the visualization panel.")


    def get_graph_data(self):
        # Extract data from the figures for saving to MATLAB
        data = {}
        for idx, fig in enumerate(self.figures):
            # Assuming your data is plotted as line plots
            ax = fig.axes[0]
            for line in ax.get_lines():
                label = line.get_label()
                if label not in ['_nolegend_']:
                    x_data, y_data = line.get_data()
                    data[f'figure_{idx}_{label}_x'] = x_data
                    data[f'figure_{idx}_{label}_y'] = y_data
        return data

    #----------------------Start new part: ----------------------
    def analyze_loss_graph(self, train_loss, val_loss, epochs):
        # Analyze the trend of the loss values over epochs
        final_train_loss = train_loss[-1]
        final_val_loss = val_loss[-1]
        analysis_text = (
            f"The training loss starts at {train_loss[0]:.4f} and ends at {final_train_loss:.4f} over {len(epochs)} epochs, "
            f"while the validation loss starts at {val_loss[0]:.4f} and ends at {final_val_loss:.4f}. "
            "This suggests that the model is learning from the training data. "
            f"{'However, ' if final_val_loss > val_loss[0] else ''}The validation loss "
            f"{'increased' if final_val_loss > val_loss[0] else 'decreased'} over time, "
            f"which {'may indicate overfitting.' if final_val_loss > val_loss[0] else 'indicates good generalization.'}"
        )
        return analysis_text

    def analyze_prediction_accuracy(self, y_true, y_pred):
        # Perform linear regression analysis for predicted vs actual reliability
        slope, intercept, r_value, p_value, std_err = linregress(y_true, y_pred)
        analysis_text = (
            f"The model's predictions have a correlation coefficient (R) of {r_value:.2f}, "
            f"indicating {'a strong' if abs(r_value) > 0.5 else 'a weak'} linear relationship with the true values. "
            f"The coefficient of determination (R^2) is {r_value**2:.3f}, "
            f"which measures the proportion of variance in the dependent variable that is predictable from the independent variable.\n"
        )
        return analysis_text

    def analyze_feature_impact(self, x_data, y_data, feature_name):
        # Simple correlation for feature impact
        correlation = np.corrcoef(x_data, y_data)[0, 1]
        analysis_text = (
            f"The impact of feature '{feature_name}' on reliability shows a correlation of {correlation:.2f}. "
            f"This {'suggests' if abs(correlation) > 0.5 else 'does not suggest'} a strong linear relationship."
        )
        return analysis_text

    #------------------End new part  !  !  ! -------------------

    def save_graphs_for_matlab(self, file_path):
        # Save the graph data to a MATLAB .mat file
        data = self.get_graph_data()
        logger.debug("🐛data is:", data)  # Debug print, you may want to remove this after fixing the issue
        scipy.io.savemat(file_path, data)



    def save_graphs_as_image(self, file_path, fig):
        # Save the graph as an image
        fig.savefig(file_path)
    
    
    def save_graphs_interactive(self):
        # Save the current figure for interactive use
        file_path = filedialog.asksaveasfilename(title="Save Interactive Graph", 
                                                filetypes=[("All Files", "*.*")],
                                                defaultextension="")
        if file_path:
            # You can use a different format like .pickle, .mat, etc.
            pass